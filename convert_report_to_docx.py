"""
Convert report.md to report.docx

Reads the generated report.md and produces a properly formatted Word document
with headings, bold text, and paragraph styles.

Usage:
    python convert_report_to_docx.py
"""

import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH


def parse_markdown_to_docx(md_path: str, docx_path: str) -> None:
    """Parse report.md and write a formatted .docx file."""
    md_text = Path(md_path).read_text(encoding='utf-8')
    lines = md_text.split('\n')

    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    # Set margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)

    i = 0
    while i < len(lines):
        line = lines[i]

        # Skip empty lines
        if line.strip() == '':
            i += 1
            continue

        # Heading 1: # Title
        if line.startswith('# ') and not line.startswith('## '):
            title = line[2:].strip()
            heading = doc.add_heading(title, level=0)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            i += 1
            continue

        # Heading 2: ## Section
        if line.startswith('## ') and not line.startswith('### '):
            title = line[3:].strip()
            doc.add_heading(title, level=1)
            i += 1
            continue

        # Heading 3: ### Subsection
        if line.startswith('### '):
            title = line[4:].strip()
            doc.add_heading(title, level=2)
            i += 1
            continue

        # Numbered list items (e.g., "1. **text**: description")
        if re.match(r'^\d+\.', line.strip()):
            paragraph = doc.add_paragraph(style='List Number')
            _add_formatted_text(paragraph, line.strip())
            i += 1
            continue

        # Regular paragraph — collect consecutive non-empty, non-heading lines
        paragraph_lines = []
        while i < len(lines):
            l = lines[i]
            if l.strip() == '':
                break
            if l.startswith('#'):
                break
            if re.match(r'^\d+\.', l.strip()) and not paragraph_lines:
                break
            paragraph_lines.append(l)
            i += 1

        if paragraph_lines:
            full_text = ' '.join(paragraph_lines)
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.first_line_indent = Cm(1.27)
            _add_formatted_text(paragraph, full_text)
            continue

        i += 1

    doc.save(docx_path)
    print(f'Report berhasil dikonversi ke: {docx_path}')


def _add_formatted_text(paragraph, text: str) -> None:
    """Add text to a paragraph with bold and italic markdown formatting."""
    # Pattern to match **bold**, *italic*, or plain text segments
    pattern = r'(\*\*.*?\*\*|\*.*?\*|[^*]+)'
    parts = re.findall(pattern, text)

    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('*') and part.endswith('*'):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        else:
            paragraph.add_run(part)


if __name__ == '__main__':
    project_root = Path(__file__).parent
    md_path = project_root / 'report.md'
    docx_path = project_root / 'report.docx'

    if not md_path.exists():
        print('Error: report.md not found. Run the notebook first to generate it.')
        raise SystemExit(1)

    parse_markdown_to_docx(str(md_path), str(docx_path))
